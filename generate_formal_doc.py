import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    """Sets cell margins (padding) in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_left_border(cell, color_hex="1F4E79", size=36):
    """Applies a thick left border and removes top, bottom, and right borders."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    # Left border
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(size))  # 36 = 4.5pt
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), color_hex)
    tcBorders.append(left)
    
    # Clear others
    for border_name in ['top', 'bottom', 'right']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
        
    tcPr.append(tcBorders)

def set_table_col_widths(table, widths):
    """Forces exact widths on a table's columns and cells."""
    for i, col in enumerate(table.columns):
        col.width = widths[i]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

def add_heading_with_spacing(doc, text, level, space_before=16, space_after=6):
    heading = doc.add_heading(text, level=level)
    p_format = heading.paragraph_format
    p_format.space_before = Pt(space_before)
    p_format.space_after = Pt(space_after)
    p_format.keep_with_next = True
    
    # Force Calibri Light / Calibri for headings with specific dark colors
    for r in heading.runs:
        r.font.name = 'Calibri Light' if level == 1 else 'Calibri'
        if level == 1:
            r.font.size = Pt(18)
            r.font.bold = True
            r.font.color.rgb = RGBColor(31, 78, 121)  # Deep Navy Blue
        elif level == 2:
            r.font.size = Pt(13.5)
            r.font.bold = True
            r.font.color.rgb = RGBColor(91, 155, 213)  # Slate Blue
        else:
            r.font.size = Pt(11.5)
            r.font.bold = True
            r.font.color.rgb = RGBColor(40, 40, 40)
    return heading

def add_page_number(run):
    """Appends a dynamic Word page number field to a run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def style_text_paragraph(p, text="", space_after=6, line_spacing=1.15):
    """Sets consistent body styling for a paragraph."""
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(38, 38, 38)  # Charcoal
        return run

def main():
    doc = Document()
    
    # --- Global Font Styles Setup ---
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(38, 38, 38)

    # --- Page Margins ---
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.different_first_page_header_footer = True
        
        # Header setup
        header = section.header
        p_hdr = header.paragraphs[0]
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_hdr_format = p_hdr.paragraph_format
        p_hdr_format.space_after = Pt(12)
        r_hdr = p_hdr.add_run("DeepCytes DNS Agent  |  Technical Architecture Documentation")
        r_hdr.font.name = 'Calibri'
        r_hdr.font.size = Pt(8.5)
        r_hdr.font.color.rgb = RGBColor(128, 128, 128)
        
        # Footer setup
        footer = section.footer
        p_ftr = footer.paragraphs[0]
        p_ftr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_ftr_format = p_ftr.paragraph_format
        p_ftr_format.space_before = Pt(12)
        
        # Add page label and page number
        r_ftr_label = p_ftr.add_run("CONFIDENTIAL  •  ")
        r_ftr_label.font.name = 'Calibri'
        r_ftr_label.font.size = Pt(8.5)
        r_ftr_label.font.color.rgb = RGBColor(128, 128, 128)
        
        r_ftr_page = p_ftr.add_run("Page ")
        r_ftr_page.font.name = 'Calibri'
        r_ftr_page.font.size = Pt(8.5)
        r_ftr_page.font.color.rgb = RGBColor(128, 128, 128)
        add_page_number(r_ftr_page)

    # --- COVER PAGE ---
    # Top spacing
    for _ in range(3):
        doc.add_paragraph()
        
    # Title Box (Single-cell table for design layout)
    cover_table = doc.add_table(rows=1, cols=1)
    cover_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = cover_table.rows[0].cells[0]
    cell.width = Inches(6.5)
    
    # Left border and background shading for a premium visual design
    set_cell_left_border(cell, color_hex="1F4E79", size=48)  # 6pt thick border
    set_cell_background(cell, "F4F6F9")                      # Off-white background
    set_cell_margins(cell, top=288, bottom=288, left=360, right=360) # Generous padding (approx 15-20pt)
    
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(8)
    run_title = p_title.add_run("DEEPCYTES DNS SECURITY AGENT")
    run_title.font.name = 'Calibri Light'
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(31, 78, 121)
    
    p_sub = cell.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("Technical Architecture and System Design Manual\nFocusing on Directory Layout, Library Dependencies, and Technical Limitations")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(100, 100, 100)
    
    p_meta = cell.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(0)
    p_meta.paragraph_format.line_spacing = 1.2
    run_meta = p_meta.add_run(
        "Prepared by: DeepCytes SOC Engineering Team\n"
        "Document Type: Formal System Architecture Documentation\n"
        "Version: 1.0 (Audit Release)\n"
        "Date: June 8, 2026"
    )
    run_meta.font.name = 'Calibri'
    run_meta.font.size = Pt(9.5)
    run_meta.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_page_break()

    # --- SECTION 1: EXECUTIVE SUMMARY ---
    add_heading_with_spacing(doc, "1. Executive Summary", level=1, space_before=12)
    
    p = doc.add_paragraph()
    style_text_paragraph(p, 
        "The DeepCytes DNS Security Agent is an enterprise-grade, lightweight, host-based passive telemetry collector "
        "and correlation engine written in Python. The agent sniffs local DNS transactions, traces their originating software process "
        "lineage on the endpoint, enriches events using multiple OSINT intelligence feeds (WHOIS, GeoIP, and Local Threat Feeds), "
        "and evaluates queries using statistical heuristics. This design allows the agent to automatically identify DNS Tunneling, "
        "Domain Generation Algorithms (DGAs), brand typosquatting, and fast-flux command-and-control behavior in real-time."
    )
    
    p2 = doc.add_paragraph()
    style_text_paragraph(p2, 
        "All events are correlated, contextualized with host and user metadata, and exported into a unified, structured JSON-Lines "
        "format (NDJSON) designed for seamless ingestion by Security Information and Event Management (SIEM) and Extended Detection and "
        "Response (XDR) data pipelines (such as Wazuh, Elastic, or Splunk). By tying network-level queries to endpoint execution states, "
        "it bridges the gap between traditional network sensors and endpoint agents."
    )

    doc.add_paragraph().paragraph_format.space_before = Pt(8)

    # --- SECTION 2: DIRECTORY STRUCTURE & FILE LAYOUT ---
    add_heading_with_spacing(doc, "2. Directory Structure and Component Layout", level=1)
    
    p = doc.add_paragraph()
    style_text_paragraph(p, 
        "The project workspace is structured modularly to separate statistical utilities, raw collectors, enrichers, and the central detection engine. "
        "This ensures proper code boundaries and maintainability:"
    )

    # Add directory structure representation in a code block
    dir_structure = (
        "DNS-BASED-EVENTS-CODES/\n"
        "│\n"
        "├── config.py                 # Global configurations, thresholds, static threat feeds, and asset profiles\n"
        "├── main.py                   # Orchestrates capture threads, handles file logs, and renders Live TUI Console\n"
        "│\n"
        "├── utils/                    # Statistical & mathematical modules\n"
        "│   ├── __init__.py           # Utility initialization\n"
        "│   ├── entropy.py            # Calculates Shannon Entropy of strings (randomness analyzer)\n"
        "│   ├── typosquat.py          # Detects brand lookalike domains using Levenshtein distance & character substitution\n"
        "│   └── dga.py                # Computes DGA confidence scores via character transition distributions\n"
        "│\n"
        "├── collectors/               # Live endpoint and network telemetry collectors\n"
        "│   ├── __init__.py           # Collector package initialization\n"
        "│   ├── dns_sniffer.py        # Captures port 53 (Scapy) and maps UDP sockets to local PIDs (with fallback simulator)\n"
        "│   └── process_collector.py  # Gathers command lines, parent lineages, binary hashes, and publisher certificates via psutil\n"
        "│\n"
        "├── enrichers/                # Intelligence, geographical, and contextual enrichment modules\n"
        "│   ├── __init__.py           # Enricher package initialization\n"
        "│   ├── whois_enricher.py     # Queries WHOIS registration details (incorporates local SQLite/JSON caching)\n"
        "│   ├── geoip_enricher.py     # Resolves response IPs to country/city/ASN (filters out private RFC 1918 IPs)\n"
        "│   ├── threat_intel.py       # Cross-references active indicators against local/remote Threat Feeds\n"
        "│   └── asset_enricher.py     # Appends hostname, OS architecture, department profiles, and user authority\n"
        "│\n"
        "└── engine/                   # Analytics, historical tracking, and correlation engines\n"
        "    ├── __init__.py           # Engine initialization\n"
        "    ├── tunneling_detector.py # Evaluates subdomain sizes, query rates, unique subdomains, and exfiltration vectors\n"
        "    ├── historical_tracker.py # Manages persistent thread-safe temporal databases (baselines first/last seen frequencies)\n"
        "    └── correlation.py        # Schema Builder: Core coordinator mapping all telemetry layers to a unified JSON document"
    )

    table_dir = doc.add_table(rows=1, cols=1)
    table_dir.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_dir = table_dir.rows[0].cells[0]
    cell_dir.width = Inches(6.5)
    set_cell_background(cell_dir, "F4F6F8")
    set_cell_margins(cell_dir, top=120, bottom=120, left=150, right=150)
    
    # Border
    tcPr = cell_dir._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for b_name in ['top', 'bottom', 'left', 'right']:
        b = OxmlElement(f'w:{b_name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')  # thin border
        b.set(qn('w:color'), 'D0D5DD')
        tcBorders.append(b)
    tcPr.append(tcBorders)

    p_dir = cell_dir.paragraphs[0]
    p_dir.paragraph_format.space_before = Pt(4)
    p_dir.paragraph_format.space_after = Pt(4)
    p_dir.paragraph_format.line_spacing = 1.05
    run_dir = p_dir.add_run(dir_structure)
    run_dir.font.name = 'Consolas'
    run_dir.font.size = Pt(8.5)
    run_dir.font.color.rgb = RGBColor(60, 60, 60)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Component layout table
    add_heading_with_spacing(doc, "2.1 Module Component Explanations", level=2)
    
    p = doc.add_paragraph()
    style_text_paragraph(p, "The following table details the classification layer and specific responsibility of each file:")

    table_comp = doc.add_table(rows=1, cols=3)
    table_comp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_comp.style = 'Table Grid'
    
    # Configure headers
    hdr_comp = table_comp.rows[0].cells
    hdr_comp[0].text = 'File / Module Path'
    hdr_comp[1].text = 'Layer Classification'
    hdr_comp[2].text = 'Functional Responsibility'
    
    for cell in hdr_comp:
        set_cell_background(cell, "1F4E79")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(255, 255, 255)

    comp_info = [
        ("config.py", "Global Settings", "Centralizes static thresholds, popular brands, local threat intelligence feeds, and corporate asset context profile."),
        ("main.py", "Core Execution", "Orchestrates multi-threaded loops for logging, simulation/sniffing, and draws the Live TUI Console Dashboard using Rich."),
        ("utils/entropy.py", "Statistical Helper", "Calculates Shannon Entropy to measure randomness in subdomains (useful for DGA and Tunneling flags)."),
        ("utils/typosquat.py", "Statistical Helper", "Applies Levenshtein edit distance and homoglyph substitution mapping to detect brand typosquatting phishing domains."),
        ("utils/dga.py", "Statistical Helper", "Computes a DGA score by checking consonant-vowel transitions and digit ratios in domain strings."),
        ("collectors/dns_sniffer.py", "Live Sniffer / Telemetry", "Sniffs port 53 UDP traffic using Scapy and queries net connections via psutil to map network transactions to local PIDs."),
        ("collectors/process_collector.py", "Live Sniffer / Telemetry", "Inspects PIDs to resolve executable paths, full CLI commands, parent process lineage, and hashes running binaries."),
        ("enrichers/whois_enricher.py", "Enrichment Layer", "Queries WHOIS to resolve domain age and registrar. Implements a local disk cache to bypass server rate limits."),
        ("enrichers/geoip_enricher.py", "Enrichment Layer", "Resolves public response IPs to country/city/ASN via ip-api.com. Bypasses lookups for RFC 1918 private IPs."),
        ("enrichers/threat_intel.py", "Enrichment Layer", "Compares domain queries and IP answers against malware/phishing indicators of compromise."),
        ("enrichers/asset_enricher.py", "Enrichment Layer", "Binds endpoint metadata (operating system, host vulnerability level, device role) to the event schema."),
        ("engine/tunneling_detector.py", "Correlation / Detection", "Measures subdomain sizes, query frequencies, and unique subdomain distributions to identify covert tunnels."),
        ("engine/historical_tracker.py", "Correlation / Detection", "Tracks temporal frequency, first seen, and last seen baselines in a persistent database."),
        ("engine/correlation.py", "Correlation / Detection", "The core Schema Builder that sequences collectors, enrichers, and detectors to compile the final JSON event.")
    ]

    for idx, (path, layer, desc) in enumerate(comp_info, 1):
        row = table_comp.add_row().cells
        row[0].text = path
        row[1].text = layer
        row[2].text = desc
        
        # Apply zebra striping
        bg_color = "F9FBFD" if idx % 2 == 0 else "FFFFFF"
        
        for i, cell in enumerate(row):
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = RGBColor(50, 50, 50)
            # Bold paths
            if i == 0:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(31, 78, 121)

    # Set Widths: Path = 1.7 in, Layer = 1.3 in, Desc = 3.5 in (total 6.5 in)
    set_table_col_widths(table_comp, [Inches(1.7), Inches(1.3), Inches(3.5)])

    doc.add_page_break()

    # --- SECTION 3: LIBRARIES USED & JUSTIFICATIONS ---
    add_heading_with_spacing(doc, "3. Library Dependencies and Justifications", level=1)
    
    p = doc.add_paragraph()
    style_text_paragraph(p, 
        "To achieve enterprise-grade telemetry collection, statistical domain analysis, and visual reporting, "
        "the DeepCytes DNS Agent leverages a hybrid stack of third-party security libraries and robust Python standard modules. "
        "The following section lists these dependencies and why they were chosen:"
    )

    table_lib = doc.add_table(rows=1, cols=4)
    table_lib.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_lib.style = 'Table Grid'
    
    # Configure headers
    hdr_lib = table_lib.rows[0].cells
    hdr_lib[0].text = 'Library'
    hdr_lib[1].text = 'Type'
    hdr_lib[2].text = 'Primary Responsibility'
    hdr_lib[3].text = 'Justification / Purpose'
    
    for cell in hdr_lib:
        set_cell_background(cell, "1F4E79")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(255, 255, 255)

    lib_info = [
        ("scapy", "Third-Party", "Raw Packet Sniffing", 
         "Enables low-level raw network interface sniffing on UDP port 53. Captures and dissects DNS protocol layers (DNSQR, DNSRR) to extract queries, responses, TTLs, and flags without needing heavy external packet brokers."),
         
        ("psutil", "Third-Party", "Socket and Process Map", 
         "Provides cross-platform APIs to query running processes, parent lineages, file paths, command lines, and network connection tables (`net_connections`). This is crucial for correlating a specific network UDP source port back to the initiating local PID."),
         
        ("python-whois", "Third-Party", "OSINT Domain Age", 
         "Interacts with public WHOIS registry servers to extract registration dates, creation timestamps, and registrars. Used to calculate domain age and flag 'newly registered domains' which are characteristic of fresh malicious infrastructure."),
         
        ("tldextract", "Third-Party", "Domain Parsing", 
         "Splits fully qualified domains (FQDNs) into subdomains, second-level domains (SLD), and top-level domains (TLD). Relies on the Public Suffix List to handle tricky domains (e.g., '.co.uk'), preventing errors in DGA, typosquatting, and tunneling algorithms."),
         
        ("rich", "Third-Party", "Live Console Dashboard", 
         "Draws an interactive, high-performance Terminal User Interface (TUI) with scroll grids, progress metrics, status bars, and formatted exception logs to provide real-time visual situational awareness for security analysts."),
         
        ("python-docx", "Third-Party", "Documentation Generation", 
         "Generates and formats Microsoft Word (.docx) technical document files programmatically to support corporate reporting, security audit compliance, and system documentation exports."),
         
        ("ipaddress", "Standard", "Network Range Check", 
         "Validates IP address types. Checks if resolved IP addresses belong to loopback (127.0.0.0/8) or private RFC 1918 ranges (10.0.0.0/8, 172.16.0.0/12, 192.168.0.0/16). Bypasses public GeoIP lookups for local IPs to prevent API limit exhaustion."),
         
        ("hashlib", "Standard", "Binary Cryptographic Hash", 
         "Computes MD5 and SHA-256 hashes of running process executables on the disk. This allows the correlation engine to append file hashes to process logs, facilitating instant comparison against threat intelligence hash databases (e.g. VirusTotal)."),
         
        ("urllib.request", "Standard", "HTTP REST Client", 
         "Makes lightweight, synchronous HTTP GET requests to the external GeoIP API (`ip-api.com`) to query geolocation metrics without introducing the external `requests` dependency, keeping the agent footprint minimal."),
         
        ("threading", "Standard", "Multithreaded Orchestration", 
         "Manages execution of concurrent operations: the packet capture thread, the background logger thread, and the foreground console TUI drawing thread run simultaneously without locking user interaction."),
         
        ("json", "Standard", "Serialization / Database", 
         "Parses global settings, serializes the complex multi-layer nested JSON events, and manages the persistent local cache files on disk (whois_cache.json, geoip_cache.json, historical_context.json)."),
         
        ("math", "Standard", "Entropy Calculation", 
         "Provides base-2 logarithmic calculations to compute the Shannon Entropy of subdomains. High entropy indicates random text, a primary heuristic for DNS tunneling payloads and DGA domains."),
         
        ("re", "Standard", "Pattern Matching", 
         "Applies regular expressions to strip domains, count consecutive character types (e.g. vowel-to-consonant ratios, consecutive digits), and analyze subdomain lengths for anomaly detection.")
    ]

    for idx, (name, lib_type, resp, just) in enumerate(lib_info, 1):
        row = table_lib.add_row().cells
        bg_color = "F9FBFD" if idx % 2 == 0 else "FFFFFF"
        
        for i, cell in enumerate(row):
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            
            p = cell.paragraphs[0]
            p.text = [name, lib_type, resp, just][i]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(50, 50, 50)
            if i == 0:
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(31, 78, 121)

    # Set Widths: Name = 1.1 in, Type = 1.0 in, Resp = 1.7 in, Just = 2.7 in (total 6.5 in)
    set_table_col_widths(table_lib, [Inches(1.1), Inches(1.0), Inches(1.7), Inches(2.7)])

    doc.add_page_break()

    # --- SECTION 4: TECHNICAL LIMITATIONS ---
    add_heading_with_spacing(doc, "4. Technical Limitations & Engineering Challenges", level=1)
    
    p = doc.add_paragraph()
    style_text_paragraph(p, 
        "During the development and testing of the DeepCytes DNS Security Agent, several fundamental technical limits, "
        "operating system constraints, and architectural challenges were identified. These points represent areas where "
        "direct implementation was restricted, along with descriptions of the root causes and mitigation strategies:"
    )

    limitations = [
        {
            "num": "4.1",
            "title": "Windows DNS Client Service (svchost.exe) Redirection Constraint",
            "challenge": "Under modern Windows operating systems, standard user applications (such as Google Chrome, PowerShell, "
                         "or malware binaries) do not resolve DNS queries by establishing direct outbound sockets to UDP Port 53. "
                         "Instead, they invoke system APIs (like GetAddrInfo), which delegate DNS resolution to the system-wide DNS Client "
                         "service (dnscache) running inside svchost.exe. As a result, when the Scapy sniffer captures a UDP Port 53 packet, "
                         "the local UDP source port mapping (via psutil.net_connections) always attributes the socket back to the "
                         "svchost.exe process ID (PID) rather than the actual originating application.",
            "impact": "Passive network sniffing on the host cannot map live DNS transactions back to their true user-space process "
                      "ID (like chrome.exe or cmd.exe) in real-time under Windows.",
            "workaround": "To bypass this, we implemented simulated correlation flows where the PID mappings are directly injected to "
                          "demonstrate Wazuh/XDR schema layouts. For true live resolution, the agent would need to trace Event Tracing for "
                          "Windows (ETW) logs (specifically Microsoft-Windows-DNS-Client, Event ID 3006) or load a custom NDIS filter driver "
                          "(like WinDivert or Sysmon) to catch DNS calls inside the kernel space."
        },
        {
            "num": "4.2",
            "title": "Administrative Privilege Constraints for Raw Sniffing",
            "challenge": "Sniffing raw network interface packets (UDP Port 53) using Scapy requires administrative privileges (root on Linux, "
                         "Administrator on Windows) to create raw network sockets.",
            "impact": "If the agent is executed in a standard user context, Scapy throws a socket creation error, causing the application to fail to sniff traffic.",
            "workaround": "We engineered an automatic fallback mechanism: if a socket error is encountered during sniffer startup, the "
                          "agent catches the exception, logs a warning, and automatically switches config.SIMULATION_MODE to True. "
                          "This falls back to simulated events that still link to real local process PIDs on the system, keeping the agent "
                          "running and visual dashboards active without crashing."
        },
        {
            "num": "4.3",
            "title": "Blindness to Encrypted DNS Protocols (DoH / DoT)",
            "challenge": "Modern operating systems and web browsers are shifting towards Encrypted DNS by default, specifically "
                         "DNS-over-HTTPS (DoH, TCP port 443) and DNS-over-TLS (DoT, TCP port 853). These protocols wrap DNS transactions "
                         "inside standard TLS encryption tunnels to secure resolvers (e.g. Cloudflare, Google DNS).",
            "impact": "Passive packet sniffing on UDP Port 53 is completely blind to encrypted DoH and DoT traffic. The agent cannot capture, "
                      "decode, or analyze these queries.",
            "workaround": "To capture these, the agent would need to act as an active local proxy/root CA provider to perform SSL inspection, "
                          "or read browser process memory, or consume internal OS log events. These techniques significantly increase the agent's "
                          "performance overhead, complexity, and security footprint."
        },
        {
            "num": "4.4",
            "title": "Third-Party API Rate Limiting and Network Latency",
            "challenge": "OSINT lookups for GeoIP (ip-api.com) and WHOIS queries are subject to strict rate limits. For instance, the free tier of "
                         "ip-api.com limits requests to 45 per minute. In addition, synchronous network socket calls to these servers block the "
                         "processing thread, causing latency spikes.",
            "impact": "Under moderate network traffic, lookup queries get throttled (returning errors), and the event processing loop slows down, "
                      "risking packet drops at the sniffer layer.",
            "workaround": "We mitigated this by: (1) checking for RFC 1918 private ranges locally to bypass network queries, (2) implementing a "
                          "local persistent JSON cache on disk that stores resolved results for 7 days, and (3) adding mock databases for common "
                          "simulated domains to guarantee instant responses during testing. For enterprise scaling, the agent would need local "
                          "databases (e.g. MaxMind GeoIP MMDB files) rather than remote web APIs."
        },
        {
            "num": "4.5",
            "title": "Passive Monitoring vs. Active Enforcement",
            "challenge": "The DeepCytes DNS Security Agent is designed as a passive telemetry collector and detection engine. It operates "
                         "by sniffing traffic or receiving events asynchronously.",
            "impact": "The agent acts purely as an Intrusion Detection System (IDS) and is unable to block, sinkhole, or terminate malicious "
                      "DNS requests (e.g., blocking DGA or typosquatted phishing links) in real-time.",
            "workaround": "Implementing active prevention requires the agent to function as a local DNS resolver/forwarder (altering system DNS settings "
                          "to point to localhost) or to register a driver with the host firewall (Windows Filtering Platform or iptables) to drop packets, "
                          "which increases the risk of system instability and deployment friction."
        },
        {
            "num": "4.6",
            "title": "Reconstruction of Tunneling Payload Data",
            "challenge": "While the agent utilizes statistics (Shannon Entropy, payload bytes, query frequency) to detect DNS tunneling, "
                         "it does not attempt to reconstruct the actual exfiltrated data stream.",
            "impact": "The system raises an alert flagging a covert channel but cannot show the analyst the exact file contents, passwords, "
                      "or shell commands that were exfiltrated inside the TXT/CNAME records.",
            "workaround": "Reconstruction requires session buffering, sorting out sequence numbers (since UDP packets can arrive out of order), "
                          "and decrypting custom payloads (since attackers often encrypt or compress exfiltrated data), which requires a "
                          "dedicated deep packet analysis (DPI) engine."
        },
        {
            "num": "4.7",
            "title": "Python CPU Bottlenecks (GIL) and I/O Thread Blocks",
            "challenge": "Calculating CPU-intensive mathematical statistics (Shannon Entropy, Levenshtein edit distance, DGA letter distributions) "
                         "for hundreds of transactions per second encounters CPU bottlenecks in Python due to the Global Interpreter Lock (GIL). "
                         "Furthermore, writing correlated events to dns_soc_events.json and updating cache files is bound by synchronous disk I/O.",
            "impact": "Under high network load (thousands of queries/sec), the agent would saturate a CPU core and begin dropping network packets.",
            "workaround": "To support high-throughput environments, a production version of the agent should be compiled in a system-level language "
                          "such as Rust or Go, or leverage asynchronous processing queues (using asyncio or multiprocessing) and a ring-buffer "
                          "structure to log events asynchronously."
        }
    ]

    for limit in limitations:
        # Title of Limitation
        add_heading_with_spacing(doc, f"{limit['num']} {limit['title']}", level=2)
        
        # Single-cell callout table to present details neatly
        card_table = doc.add_table(rows=1, cols=1)
        card_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        card_table.autofit = False
        
        cell = card_table.rows[0].cells[0]
        cell.width = Inches(6.5)
        set_cell_background(cell, "F8F9FA")
        set_cell_left_border(cell, color_hex="5B9BD5", size=24) # Muted blue border
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        
        # Challenge paragraph
        p_ch = cell.paragraphs[0]
        p_ch.paragraph_format.space_after = Pt(4)
        p_ch.paragraph_format.line_spacing = 1.15
        run_lbl1 = p_ch.add_run("Description of Challenge: ")
        run_lbl1.font.name = 'Calibri'
        run_lbl1.font.bold = True
        run_lbl1.font.size = Pt(9.5)
        run_lbl1.font.color.rgb = RGBColor(40, 40, 40)
        
        run_val1 = p_ch.add_run(limit["challenge"])
        run_val1.font.name = 'Calibri'
        run_val1.font.size = Pt(9.5)
        run_val1.font.color.rgb = RGBColor(60, 60, 60)
        
        # Impact paragraph
        p_imp = cell.add_paragraph()
        p_imp.paragraph_format.space_after = Pt(4)
        p_imp.paragraph_format.line_spacing = 1.15
        run_lbl2 = p_imp.add_run("Security Impact: ")
        run_lbl2.font.name = 'Calibri'
        run_lbl2.font.bold = True
        run_lbl2.font.size = Pt(9.5)
        run_lbl2.font.color.rgb = RGBColor(180, 50, 50) # Crimson for impact
        
        run_val2 = p_imp.add_run(limit["impact"])
        run_val2.font.name = 'Calibri'
        run_val2.font.size = Pt(9.5)
        run_val2.font.color.rgb = RGBColor(60, 60, 60)
        
        # Workaround paragraph
        p_work = cell.add_paragraph()
        p_work.paragraph_format.space_after = Pt(0)
        p_work.paragraph_format.line_spacing = 1.15
        run_lbl3 = p_work.add_run("Mitigation / Workaround: ")
        run_lbl3.font.name = 'Calibri'
        run_lbl3.font.bold = True
        run_lbl3.font.size = Pt(9.5)
        run_lbl3.font.color.rgb = RGBColor(31, 78, 121) # Deep navy for mitigation
        
        run_val3 = p_work.add_run(limit["workaround"])
        run_val3.font.name = 'Calibri'
        run_val3.font.size = Pt(9.5)
        run_val3.font.color.rgb = RGBColor(60, 60, 60)

        # Padding space after the card
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(4)
        p_space.paragraph_format.space_after = Pt(4)

    # Save Document
    doc_path = "DeepCytes_DNS_Agent_Formal_Documentation.docx"
    doc.save(doc_path)
    print(f"Formal Technical Documentation generated successfully: {os.path.abspath(doc_path)}")

if __name__ == "__main__":
    main()


