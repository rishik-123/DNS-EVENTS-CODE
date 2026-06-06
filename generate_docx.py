import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_with_spacing(doc, text, level, space_before=12, space_after=6):
    heading = doc.add_heading(text, level=level)
    p_format = heading.paragraph_format
    p_format.space_before = Pt(space_before)
    p_format.space_after = Pt(space_after)
    p_format.keep_with_next = True
    return heading

def main():
    doc = Document()
    
    # --- Page Margins ---
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- Title Page / Header ---
    title = doc.add_paragraph()
    title_format = title.paragraph_format
    title_format.space_before = Pt(24)
    title_format.space_after = Pt(6)
    title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("DeepCytes DNS Security Agent")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(31, 78, 121)  # Deep Navy Blue

    subtitle = doc.add_paragraph()
    subtitle_format = subtitle.paragraph_format
    subtitle_format.space_after = Pt(36)
    subtitle_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run("Mini SOC Telemetry Collector & Correlation Engine\nProject Documentation and Design Manual")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph("─" * 60).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Section 1: Executive Summary ---
    add_heading_with_spacing(doc, "1. Executive Summary", level=1, space_before=24)
    p1 = doc.add_paragraph(
        "The DeepCytes DNS Security Agent is a modular, lightweight host-based security telemetry collector designed for "
        "modern Security Operations Centers (SOCs). Written in Python, the agent sniffs local DNS transactions, traces their "
        "originating software process lineages on the endpoint, enriches events using multiple OSINT intelligence feeds (WHOIS, GeoIP, "
        "Threat Feeds), and evaluates queries via statistical heuristics to automatically identify DNS Tunneling, "
        "Domain Generation Algorithms (DGAs), typosquatting, and fast-flux command-and-control behavior."
    )
    p1.paragraph_format.line_spacing = 1.15
    p1.paragraph_format.space_after = Pt(12)

    # --- Section 2: Directory and File Layout ---
    add_heading_with_spacing(doc, "2. File Layout and Component Explanations", level=1)
    doc.add_paragraph(
        "The application is structured modularly. The table below outlines what each folder and file does, "
        "ensuring clean boundaries between collection, enrichment, analysis, and execution:"
    )

    # Create Table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'File / Module Path'
    hdr_cells[1].text = 'Layer Classification'
    hdr_cells[2].text = 'Functional Responsibility'
    
    # Set headers styling
    for i, cell in enumerate(hdr_cells):
        set_cell_background(cell, "1F4E79")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    files_info = [
        ("config.py", "Global Configuration", "Defines log directories, thresholds for DGA/Tunneling, popular brands list, local static threat intelligence data, and corporate asset context information."),
        ("main.py", "Agent Core & UI Engine", "Orchestrates threads. Captures raw packets or triggers simulated alerts, directs them into the correlation engine, appends events to output logs, and outputs a Live TUI Console Dashboard using the Rich library."),
        ("utils/entropy.py", "Helper Utilities", "Implements Shannon Entropy to measure character randomness. Used by DGA and DNS Tunneling detectors to flag high-randomness domains."),
        ("utils/typosquat.py", "Helper Utilities", "Uses Levenshtein edit distance and visual substitution mapping (e.g., '0' for 'o', '1' for 'i') to detect lookalike phishing domains targeting key brands."),
        ("utils/dga.py", "Helper Utilities", "Calculates a DGA confidence score between 0.0 and 1.0 by evaluating consonant-to-vowel ratios, consecutive character lengths, and numbers inside domain names."),
        ("collectors/dns_sniffer.py", "Telemetry Collector", "Filters and captures traffic on UDP port 53 (using Scapy) and traces active sockets using psutil connections to bind the request to a local Process ID (PID)."),
        ("collectors/process_collector.py", "Telemetry Collector", "Inspects a given PID to retrieve command lines, executable paths, parent processes, user privileges, code signers (Windows PE metadata), and file MD5/SHA256 hashes."),
        ("enrichers/whois_enricher.py", "Enrichment Layer", "Queries WHOIS databases to extract domain age, creation date, and registrar. Protects against server throttling using a persistent local JSON cache."),
        ("enrichers/geoip_enricher.py", "Enrichment Layer", "Resolves public response IPs to country, city, and ASN using ip-api.com. Bypasses lookup for local/private range IPs, marking them instantly as 'Internal'."),
        ("enrichers/threat_intel.py", "Enrichment Layer", "Compares domain names and IP addresses against active threat intelligence indicators (IP blocks, malware, phishing domains)."),
        ("enrichers/asset_enricher.py", "Enrichment Layer", "Binds structural machine context (device criticality, department, operating system) and identity roles (privilege level) to the SOC event."),
        ("engine/tunneling_detector.py", "Detection Engine", "Evaluates subdomain payload size, rolling query rates, and unique subdomain requests under a parent domain to identify covert communication channels."),
        ("engine/historical_tracker.py", "Detection Engine", "Manages a thread-safe historical tracker. Calculates baseline frequencies and rolling queries-per-minute values in a JSON database."),
        ("engine/correlation.py", "Correlation Engine", "Coordinates the 'Schema Builder'. Ingests raw telemetry and sequentially runs all collectors, enrichers, and detectors to compile a unified multi-layer JSON SOC event.")
    ]

    for path, layer, desc in files_info:
        row_cells = table.add_row().cells
        row_cells[0].text = path
        row_cells[1].text = layer
        row_cells[2].text = desc
        # Bold the paths
        for p in row_cells[0].paragraphs:
            for r in p.runs:
                r.font.bold = True

    # Add space after table
    doc.add_paragraph().paragraph_format.space_before = Pt(12)

    # --- Section 3: How Everything Takes Place ---
    add_heading_with_spacing(doc, "3. Event Processing and Correlation Flow", level=1)
    doc.add_paragraph(
        "When the DeepCytes DNS Agent executes, telemetry and correlation take place through a structured pipeline:"
    )
    
    flow_steps = [
        "Packet Sniffing or Simulation: The sniffer starts on a separate thread, listening on UDP port 53. If running in Simulation Mode, the agent cycles through predefined attack scenarios (Normal, Typosquatting, DGA, Tunneling, Fast Flux, Split-Brain) and grabs real process IDs currently active on the host.",
        "Local Socket Mapping: When a DNS query occurs, the collector catches the source port (sport) of the local UDP socket. It checks active connection bindings via psutil.net_connections() to extract the specific Process ID (PID) that spawned the query.",
        "Endpoint Process Inspection: The PID is sent to the process collector. It extracts the parent process name, execution command line, file hashes, and publisher signature, defining the 'Process Layer'.",
        "Domain & Statistical Analysis: The query is broken down into subdomain, SLD, and TLD. The agent calculates Shannon entropy, typosquatting status, and DGA confidence scores.",
        "Intelligence Enrichment: The resolved response IPs are queried for GeoIP (mapping country and ASN) and WHOIS (resolving domain age). Public feeds and local databases are queried for threat reputation votes.",
        "Temporal Correlation: The correlation engine updates the historical database to increment query counts, track first/last seen timestamps, and compute rolling queries-per-minute rates.",
        "Heuristic Classification: The tunneling detector merges query rates, payload sizes, and unique subdomain counts to determine if covert data exfiltration is occurring.",
        "Unified Schema Serialization: The correlation engine compiles all information into a multi-layered JSON structure. The final document is written to logs/dns_soc_events.json and displayed on the console dashboard."
    ]
    
    for i, step in enumerate(flow_steps, 1):
        p = doc.add_paragraph(style='List Bullet')
        run_bold = p.add_run(f"Step {i}: ")
        run_bold.bold = True
        p.add_run(step)
        p.paragraph_format.space_after = Pt(4)

    # --- Section 4: Why Unsolicited Domains (e.g. github.com) Appear ---
    add_heading_with_spacing(doc, "4. The Presence of Background and Unsolicited DNS Traffic", level=1)
    
    doc.add_paragraph(
        "A common point of confusion during security analysis is seeing DNS events for websites "
        "that the local analyst did not search for (such as github.com, microsoft.com, or randomized domains). "
        "There are two primary reasons why these appear in dns_soc_events.json:"
    )

    p_reason1 = doc.add_paragraph()
    run_r1 = p_reason1.add_run("Reason A: Simulation Mode Loop (Default)\n")
    run_r1.bold = True
    p_reason1.add_run(
        "By default, config.SIMULATION_MODE is set to True. In this mode, the agent simulates diverse, real-world network traffic "
        "to validate and demonstrate threat detection capabilities (DGA, Tunneling, and Phishing) without needing administrative network interfaces. "
        "As part of this simulator loop, it generates mock requests for common benign domains like github.com, gmail.com, and company portals "
        "every few seconds to serve as a telemetry baseline."
    )
    p_reason1.paragraph_format.left_indent = Inches(0.25)
    p_reason1.paragraph_format.space_after = Pt(12)

    p_reason2 = doc.add_paragraph()
    run_r2 = p_reason2.add_run("Reason B: Background Operating System & Application Telemetry\n")
    run_r2.bold = True
    p_reason2.add_run(
        "If SIMULATION_MODE is set to False, the agent starts sniffing raw network interfaces. In modern operating systems (Windows 11, Linux, macOS), "
        "hundreds of background DNS queries occur continuously without user interaction. Examples include:\n"
        "  • Update services checking for patches (e.g., windowsupdate.com, googleapis.com)\n"
        "  • Cloud synchronization tools checking connectivity (e.g., onedrive.live.com, dropbox.com)\n"
        "  • Web browsers performing background pre-fetching or certificate verification (e.g., github.com, digicert.com)\n"
        "  • Active AI agents or editors querying package repositories\n"
        "The agent captures all active traffic on UDP Port 53, bringing absolute visibility to these background operations."
    )
    p_reason2.paragraph_format.left_indent = Inches(0.25)
    p_reason2.paragraph_format.space_after = Pt(12)

    # --- Section 5: Key Edge Cases Handled ---
    add_heading_with_spacing(doc, "5. Critical Edge Cases Handled", level=1)
    
    edge_cases = [
        ("Privilege Requirements (Permission Fallback)", 
         "Sniffing UDP port 53 requires administrative or root privileges. If the agent runs without these, it catches the socket permission error and automatically redirects configuration to Simulation Mode, ensuring uninterrupted demonstration capabilities."),
        
        ("Windows DNS Client Service (svchost.exe)", 
         "Windows processes delegate DNS resolutions to the centralized DNS Client service (dnscache), making all local network sockets originate from svchost.exe. To handle this, the correlation engine supports EDR process inputs and maps simulations to their actual originating software (e.g., chrome.exe, powershell.exe)."),
        
        ("WHOIS & GeoIP Server Throttling", 
         "WHOIS and public GeoIP lookup endpoints strictly rate-limit queries. The agent implements two mitigation mechanisms: First, it short-circuits lookups for RFC 1918 and loopback IPs, identifying them locally as 'Internal Network' without using API quota. Second, it maintains persistent JSON databases on disk to cache resolved values, dramatically increasing execution speed."),
        
        ("Fast Flux Command-and-Control Networks", 
         "Malicious domains evade IP-based blocking by continuously swapping active IP associations with very low Cache TTL values. The agent parses response payloads, tracks low TTL ranges (TTL < 15), and triggers a SUSPICIOUS_LOW_TTL_FAST_FLUX alert if dynamic rotations are detected.")
    ]

    for title_ec, desc_ec in edge_cases:
        p_ec = doc.add_paragraph()
        run_title = p_ec.add_run(f"•  {title_ec}: ")
        run_title.bold = True
        p_ec.add_run(desc_ec)
        p_ec.paragraph_format.space_after = Pt(6)

    # Save Document
    doc_path = "DeepCytes_DNS_Agent_Documentation.docx"
    doc.save(doc_path)
    print(f"Documentation generated successfully: {os.path.abspath(doc_path)}")

if __name__ == "__main__":
    main()
