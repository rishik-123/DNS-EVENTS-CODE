import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_with_spacing(doc, text, level, space_before=18, space_after=6):
    heading = doc.add_heading(text, level=level)
    p_format = heading.paragraph_format
    p_format.space_before = Pt(space_before)
    p_format.space_after = Pt(space_after)
    p_format.keep_with_next = True
    return heading

def main():
    doc = Document()
    
    # --- Page Margins ---
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # --- Title Page ---
    title = doc.add_paragraph()
    title_format = title.paragraph_format
    title_format.space_before = Pt(36)
    title_format.space_after = Pt(6)
    title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("DeepCytes DNS Security Agent")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle_format = subtitle.paragraph_format
    subtitle_format.space_after = Pt(36)
    subtitle_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run("Complete Source Code Compilation")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(16)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph("─" * 60).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_intro = doc.add_paragraph(
        "This document contains the complete, unabridged source code for all modules of the DeepCytes DNS Security Agent. "
        "Each section displays the full file name, folder location, and implementation code. "
        "Code sections are styled in monospaced format inside shaded containers for clear code review."
    )
    p_intro.paragraph_format.space_after = Pt(24)

    # List of files in execution/logical order
    code_files = [
        "config.py",
        "utils/__init__.py",
        "utils/entropy.py",
        "utils/typosquat.py",
        "utils/dga.py",
        "collectors/__init__.py",
        "collectors/process_collector.py",
        "collectors/dns_sniffer.py",
        "enrichers/__init__.py",
        "enrichers/whois_enricher.py",
        "enrichers/geoip_enricher.py",
        "enrichers/threat_intel.py",
        "enrichers/asset_enricher.py",
        "engine/__init__.py",
        "engine/tunneling_detector.py",
        "engine/historical_tracker.py",
        "engine/correlation.py",
        "main.py"
    ]

    for idx, rel_path in enumerate(code_files, 1):
        full_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), rel_path)
        
        # Add heading
        heading_text = f"{idx}. File: {rel_path}"
        add_heading_with_spacing(doc, heading_text, level=1)
        
        # Read file code
        if not os.path.exists(full_path):
            doc.add_paragraph(f"[Error: File {rel_path} not found on disk.]")
            continue
            
        try:
            with open(full_path, "r", encoding="utf-8") as f:
                code_content = f.read()
        except Exception as e:
            code_content = f"[Error reading file: {e}]"

        # Create single-cell table for code styling (looks like a markdown code block)
        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.allow_autofit = False
        
        # Set block width to match page margin area (approx 6.9 inches)
        cell = table.rows[0].cells[0]
        cell.width = Inches(6.9)
        set_cell_background(cell, "F2F2F2")  # Light gray background

        # Add code to cell
        p_code = cell.paragraphs[0]
        p_code_format = p_code.paragraph_format
        p_code_format.space_before = Pt(4)
        p_code_format.space_after = Pt(4)
        p_code_format.line_spacing = 1.0
        
        run_code = p_code.add_run(code_content)
        run_code.font.name = 'Consolas'
        run_code.font.size = Pt(8.5)
        run_code.font.color.rgb = RGBColor(40, 40, 40)
        
        # Page break after code block, except for the last file
        if idx < len(code_files):
            doc.add_page_break()

    # Save
    out_path = "DeepCytes_DNS_Agent_SourceCode.docx"
    doc.save(out_path)
    print(f"Source code compilation generated successfully: {os.path.abspath(out_path)}")

if __name__ == "__main__":
    main()
