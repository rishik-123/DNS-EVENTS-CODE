import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from PIL import Image, ImageDraw, ImageFont

def draw_arrow(draw, start, end, color='#555555', width=2):
    # Draw line from start (x1, y1) to end (x2, y2)
    draw.line([start, end], fill=color, width=width)
    
    # Calculate arrowhead coordinates
    x1, y1 = start
    x2, y2 = end
    
    dx = x2 - x1
    dy = y2 - y1
    length = (dx**2 + dy**2)**0.5
    if length == 0:
        return
    
    # Normalize
    ux = dx / length
    uy = dy / length
    
    # Arrow size
    arrow_len = 10
    arrow_width = 5
    
    # Perpendicular vector
    px = -uy
    py = ux
    
    # Arrow tip is at (x2, y2)
    bx = x2 - ux * arrow_len
    by = y2 - uy * arrow_len
    
    lx = bx + px * arrow_width
    ly = by + py * arrow_width
    rx = bx - px * arrow_width
    ry = by - py * arrow_width
    
    draw.polygon([(x2, y2), (lx, ly), (rx, ry)], fill=color)

def draw_flow_chart():
    # Size: 1200 x 1400 (light background)
    w, h = 1200, 1400
    img = Image.new('RGB', (w, h), color='#FFFFFF')
    draw = ImageDraw.Draw(img)

    # Attempt to load fonts
    try:
        title_font = ImageFont.truetype("arial.ttf", 24)
        subtitle_font = ImageFont.truetype("arial.ttf", 13)
        lane_font = ImageFont.truetype("arial.ttf", 14)
        bold_font = ImageFont.truetype("arial.ttf", 12)
        regular_font = ImageFont.truetype("arial.ttf", 11)
        code_font = ImageFont.truetype("cour.ttf", 10.5)
    except IOError:
        # Fallback to default
        title_font = subtitle_font = lane_font = bold_font = regular_font = code_font = ImageFont.load_default()

    # Draw Title Header
    draw.text((50, 30), "SOC LOG TELEMETRY INGESTION PIPELINE FLOW CHART", fill="#1F4E79", font=title_font)
    draw.text((50, 65), "End-to-End Processing, Routing, Database Ingest, and Offset Commit Lifecycle", fill="#555555", font=subtitle_font)
    draw.line([(50, 95), (1150, 95)], fill="#CCCCCC", width=2)

    # Helper function to draw boxes
    def draw_box(xy, fill_color, border_color, title, lines, title_color="#1F4E79", text_color="#333333"):
        draw.rounded_rectangle(xy, radius=6, fill=fill_color, outline=border_color, width=2)
        x1, y1, x2, y2 = xy
        
        # Center title
        title_w = draw.textlength(title, font=bold_font)
        draw.text((x1 + (x2 - x1 - title_w)/2, y1 + 10), title, fill=title_color, font=bold_font)
        
        y_text = y1 + 32
        for line in lines:
            if line.startswith("•"):
                draw.text((x1 + 15, y_text), line, fill=text_color, font=regular_font)
            else:
                line_w = draw.textlength(line, font=regular_font)
                draw.text((x1 + (x2 - x1 - line_w)/2, y_text), line, fill=text_color, font=regular_font)
            y_text += 18

    # --- STAGE 1: COLLECTION (y: 110 - 320) ---
    draw.rectangle([(50, 110), (1150, 310)], fill="#F8FAFC", outline="#E2E8F0", width=1)
    draw.text((65, 120), "STAGE 1: ENDPOINT TELEMETRY COLLECTION", fill="#1F4E79", font=lane_font)

    # Box 1: 10 Log Sources on Endpoint
    draw_box(
        xy=(80, 160, 360, 270),
        fill_color="#E3F2FD",
        border_color="#1E88E5",
        title="10 Log Sources on Endpoint",
        lines=[
            "DNS Queries, Process execution,",
            "Network sockets, Registry keys,",
            "Syslog, File writes, etc.",
            "Buffered locally in memory queue"
        ]
    )

    # Box 2: JSON Envelope Wrapper
    draw_box(
        xy=(460, 160, 780, 270),
        fill_color="#E8F5E9",
        border_color="#43A047",
        title="Standardized JSON Wrapper",
        lines=[
            "Sets metadata: event_id (UUID),",
            "timestamp, host_id, log_type.",
            "• KEY: Hostname / endpoint_id",
            "• VALUE: Serialized JSON Envelope"
        ],
        title_color="#2E7D32"
    )

    # Draw arrow Box 1 -> Box 2
    draw_arrow(draw, (360, 215), (460, 215), color="#1E88E5", width=2)
    # Draw arrow Box 2 -> down to Stage 2
    draw_arrow(draw, (620, 270), (620, 350), color="#43A047", width=2)


    # --- STAGE 2: KAFKA INGESTION & PARTITIONING (y: 320 - 680) ---
    draw.rectangle([(50, 310), (1150, 670)], fill="#FFFFFF", outline="#E2E8F0", width=1)
    draw.text((65, 320), "STAGE 2: KAFKA INGESTION & ROUTING TIER", fill="#1F4E79", font=lane_font)

    # Box 3: Topic Router
    draw_box(
        xy=(470, 350, 770, 440),
        fill_color="#FFF3E0",
        border_color="#FB8C00",
        title="Kafka Topic Routing Engine",
        lines=[
            "Evaluates internal 'alerts' list array:",
            "IF empty -> Route to RAW / IF active -> ALERTS"
        ],
        title_color="#E65100"
    )

    # Two topics
    draw_box(
        xy=(150, 470, 420, 540),
        fill_color="#FFEBEE",
        border_color="#E53935",
        title="Topic: dns-alerts",
        lines=["High priority, suspicious findings."],
        title_color="#C62828"
    )

    draw_box(
        xy=(800, 470, 1070, 540),
        fill_color="#E8EAF6",
        border_color="#3949AB",
        title="Topic: dns-events-raw",
        lines=["High volume, standard resolutions."],
        title_color="#283593"
    )

    # Arrows Router -> Topics
    draw_arrow(draw, (470, 395), (285, 470), color="#FB8C00", width=2)
    draw_arrow(draw, (770, 395), (935, 470), color="#FB8C00", width=2)

    # Hashing Box
    draw_box(
        xy=(420, 560, 820, 650),
        fill_color="#F3E5F5",
        border_color="#8E24AA",
        title="Key-Based Partition Hashing Strategy",
        lines=[
            "Formula: Target Partition = Hash(endpoint_id) % 3",
            "Ensures all events of a specific endpoint map to the same partition,",
            "guaranteeing sequential in-order delivery of host activities."
        ],
        title_color="#6A1B9A"
    )

    # Arrows Topics -> Hashing Box
    draw_arrow(draw, (285, 540), (520, 560), color="#8E24AA", width=2)
    draw_arrow(draw, (935, 540), (720, 560), color="#8E24AA", width=2)

    # Partition Boxes
    p_y = 685
    draw.rounded_rectangle([(350, p_y), (480, p_y+40)], radius=4, fill="#ECEFF1", outline="#607D8B", width=1)
    draw.text((375, p_y+12), "Partition 0", fill="#37474F", font=bold_font)

    draw.rounded_rectangle([(535, p_y), (665, p_y+40)], radius=4, fill="#ECEFF1", outline="#607D8B", width=1)
    draw.text((560, p_y+12), "Partition 1", fill="#37474F", font=bold_font)

    draw.rounded_rectangle([(720, p_y), (850, p_y+40)], radius=4, fill="#ECEFF1", outline="#607D8B", width=1)
    draw.text((745, p_y+12), "Partition 2", fill="#37474F", font=bold_font)

    # Hashing Box -> Partitions
    draw_arrow(draw, (490, 650), (415, p_y), color="#8E24AA", width=2)
    draw_arrow(draw, (620, 650), (600, p_y), color="#8E24AA", width=2)
    draw_arrow(draw, (750, 650), (785, p_y), color="#8E24AA", width=2)


    # --- STAGE 3: CONSUMER PROCESSING (y: 750 - 950) ---
    draw.rectangle([(50, 750), (1150, 950)], fill="#F8FAFC", outline="#E2E8F0", width=1)
    draw.text((65, 760), "STAGE 3: CONSUMPTION, PARSING & ENRICHMENT TIER", fill="#1F4E79", font=lane_font)

    # Consumer Engine
    draw_box(
        xy=(250, 800, 950, 910),
        fill_color="#E0F7FA",
        border_color="#00ACC1",
        title="Real-Time Consumer Processing Service",
        lines=[
            "• Normalizes telemetry records into standardized Elastic Common Schema (ECS).",
            "• Enriches events: GeoIP (MaxMind RAM cache) & WHOIS (persistent JSON cache tables).",
            "• Security Checks: Matches domain/IP indicators against high-speed threat intel feeds (Redis).",
            "• Statistical Heuristics: Shannon entropy check on query strings to detect DGA/Tunneling."
        ],
        title_color="#006064"
    )

    # Partitions -> Consumer Engine
    draw_arrow(draw, (415, p_y+40), (415, 800), color="#00ACC1", width=2)
    draw_arrow(draw, (600, p_y+40), (600, 800), color="#00ACC1", width=2)
    draw_arrow(draw, (785, p_y+40), (785, 800), color="#00ACC1", width=2)


    # --- STAGE 4: STORAGE & RETENTION (y: 950 - 1380) ---
    draw.rectangle([(50, 950), (1150, 1370)], fill="#FFFFFF", outline="#E2E8F0", width=1)
    draw.text((65, 960), "STAGE 4: FAULT-TOLERANT STORAGE & RETENTION LIFECYCLE", fill="#1F4E79", font=lane_font)

    # Database Batch Insert
    draw_box(
        xy=(350, 990, 850, 1070),
        fill_color="#E8F5E9",
        border_color="#43A047",
        title="Database Batch Write",
        lines=[
            "Executes batch SQL insert into ClickHouse / TimescaleDB.",
            "Uses unique 'event_id' (UUIDv4) for deduplication (idempotence constraint)."
        ],
        title_color="#1B5E20"
    )

    # Arrow Consumer -> Database Write
    draw_arrow(draw, (600, 910), (600, 990), color="#00ACC1", width=2)

    # Manual Offset Commit Box (Left)
    draw_box(
        xy=(80, 1140, 420, 1240),
        fill_color="#E0F2F1",
        border_color="#00897B",
        title="Manual Commit Offset (ACK)",
        lines=[
            "Advance consumer cursor in Kafka.",
            "Config: enable.auto.commit = False",
            "Ensures NO messages are skipped on crash."
        ],
        title_color="#004D40"
    )

    # Tiered Retention Storage (Right)
    draw_box(
        xy=(780, 1100, 1110, 1350),
        fill_color="#ECEFF1",
        border_color="#607D8B",
        title="Tiered Storage Lifecycle",
        lines=[
            "• Hot Tier (0-15 Days):",
            "  Fast SSDs (Primary ClickHouse DB).",
            "  Sub-second search speeds.",
            "• Warm Tier (16-90 Days):",
            "  Compressed HDD storage tables.",
            "  Searchable but lower IOPs.",
            "• Cold Tier (91+ Days):",
            "  Parquet tables moved to AWS S3/MinIO.",
            "  Highly compressed, queryable serverless."
        ],
        title_color="#37474F",
        text_color="#455A64"
    )

    # Arrow Database Write -> Manual Commit
    draw_arrow(draw, (350, 1030), (250, 1140), color="#43A047", width=2)
    # Arrow Database Write -> Tiered Storage
    draw_arrow(draw, (850, 1030), (945, 1100), color="#43A047", width=2)

    # Loop back arrow from Manual Commit offset to Partitions (Stage 2)
    # Draw line left, up and connect
    draw.line([(80, 1190), (65, 1190), (65, 705), (330, 705)], fill="#00897B", width=2)
    draw_arrow(draw, (330, 705), (350, 705), color="#00897B", width=2)
    draw.text((70, 715), "Advance Offset cursor in broker", fill="#00897B", font=regular_font)

    # Save diagram
    img.save("kafka_integration_diagram.png")
    print("New visual flow chart diagram created successfully: kafka_integration_diagram.png")

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_with_spacing(doc, text, level, space_before=12, space_after=6):
    heading = doc.add_heading(text, level=level)
    p_format = heading.paragraph_format
    p_format.space_before = Pt(space_before)
    p_format.space_after = Pt(space_after)
    p_format.keep_with_next = True
    
    # Custom heading styles
    for run in heading.runs:
        run.font.name = 'Arial'
        if level == 1:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(31, 78, 121) # Deep Navy
        elif level == 2:
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 168, 150) # Teal
    return heading

def main():
    # 1. Generate the flow chart
    draw_flow_chart()

    # 2. Start DOCX creation
    doc = Document()
    
    # Page margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- Title Section ---
    title = doc.add_paragraph()
    title_format = title.paragraph_format
    title_format.space_before = Pt(30)
    title_format.space_after = Pt(6)
    title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("SOC Telemetry & Kafka Integration Design Specification")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(31, 78, 121)  # Deep Navy Blue

    subtitle = doc.add_paragraph()
    subtitle_format = subtitle.paragraph_format
    subtitle_format.space_after = Pt(24)
    subtitle_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run("A Server-Side Engineering Design Document for Multi-Source Telemetry Ingestion, Fault-Tolerant Processing, and Tiered Archival Storage")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(128, 128, 128)

    divider = doc.add_paragraph()
    divider.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider.paragraph_format.space_after = Pt(24)
    run_div = divider.add_run("─" * 70)
    run_div.font.color.rgb = RGBColor(31, 78, 121)

    # --- Section: Executive Architecture ---
    add_heading_with_spacing(doc, "System Architecture & Design Questions", level=1, space_before=18)
    
    intro_p = doc.add_paragraph(
        "This document details the architectural decisions and system design patterns for ingesting high-throughput security event logs "
        "from host endpoints into an Enterprise Security Operations Center (SOC). From a server-side engineering perspective, "
        "building a pipeline that is resilient, scalable, and highly queryable requires addressing several database, queuing, "
        "and data serialization challenges. Below are the design specifications addressing the nine core questions of this integration."
    )
    intro_p.paragraph_format.line_spacing = 1.15
    intro_p.paragraph_format.space_after = Pt(12)

    # --- Question 1 ---
    add_heading_with_spacing(doc, "1) Ingesting 10 Different Logs Collected for a Single Endpoint", level=2)
    p1 = doc.add_paragraph(
        "To successfully collect and push 10 different log types (e.g., DNS Queries, Process Creations, Network Sockets, Registry Changes, "
        "Windows Event Logs, File Operations, System Resource Metrics, Active User Sessions, Firewall Violations, and Kernel Audit Records) "
        "from a single endpoint, we design a unified host shipper/collector agent on the endpoint (e.g., Filebeat, Vector, or a custom daemon). "
        "Instead of transmitting these logs as independent, unstructured events, the agent serializes them into a unified, structured JSON metadata envelope. "
        "This envelope acts as a common contract between the endpoint and the backend server, attaching universal metadata to allow cross-log correlation:"
    )
    p1.paragraph_format.line_spacing = 1.15
    p1.paragraph_format.space_after = Pt(8)

    # Bullet points for Envelope metadata
    bullet1_1 = doc.add_paragraph(style='List Bullet')
    run = bullet1_1.add_run("Global Envelope Headers: ")
    run.bold = True
    bullet1_1.add_run("Contains universal telemetry attributes: `event_id` (UUIDv4), `timestamp` (event_time in UTC), `log_type` (e.g., DNS, PROCESS), and `endpoint_id` / `hostname`.")
    bullet1_1.paragraph_format.space_after = Pt(3)

    bullet1_2 = doc.add_paragraph(style='List Bullet')
    run = bullet1_2.add_run("Polymorphic Payload Wrapper: ")
    run.bold = True
    bullet1_2.add_run("Contains an inner `data` object whose schema varies dynamically based on the `log_type` (e.g., process command line for process logs, destination port for firewall logs).")
    bullet1_2.paragraph_format.space_after = Pt(3)

    bullet1_3 = doc.add_paragraph(style='List Bullet')
    run = bullet1_3.add_run("Sender Buffer & Batching: ")
    run.bold = True
    bullet1_3.add_run("To minimize network overhead and socket exhaustion, the agent buffers logs in local memory and sends them in compressed batches (micro-batches) to the ingestion server via TCP/TLS sockets.")
    bullet1_3.paragraph_format.space_after = Pt(12)

    # --- Question 2 ---
    add_heading_with_spacing(doc, "2) Handling Log Subtypes, Details, and Depths", level=2)
    p2 = doc.add_paragraph(
        "Each log type has unique subtypes, nested details, and varying depth levels. For example, a Process Log may have subtypes "
        "like `CREATION`, `TERMINATION`, or `INJECTION` and contain deep metadata such as code-signing certificates and execution chains. "
        "To manage this complexity without schema conflicts, we design a hierarchical polymorphic schema (e.g., JSON Schema, Avro, or Protobuf):"
    )
    p2.paragraph_format.line_spacing = 1.15
    p2.paragraph_format.space_after = Pt(8)

    bullet2_1 = doc.add_paragraph(style='List Bullet')
    run = bullet2_1.add_run("Nested Objects: ")
    run.bold = True
    bullet2_1.add_run("The root envelope defines common metadata, while subtype-specific attributes are encapsulated in nested sub-objects. "
                           "This maintains a stable root structure while supporting deep nested fields (e.g., `data.process.parent_process.hashes.sha256`).")
    bullet2_1.paragraph_format.space_after = Pt(3)

    bullet2_2 = doc.add_paragraph(style='List Bullet')
    run = bullet2_2.add_run("Taxonomy Standardization: ")
    run.bold = True
    bullet2_2.add_run("We map fields to a consistent standard like the Elastic Common Schema (ECS). "
                           "This ensures variables across different operating systems (such as `username`, `login_user`, or `uid`) map to a single "
                           "consistent database column (e.g., `user.name`).")
    bullet2_2.paragraph_format.space_after = Pt(12)

    # --- Question 3 ---
    add_heading_with_spacing(doc, "3) Managing Asynchronous and Out-of-Order Log Arrival", level=2)
    p3 = doc.add_paragraph(
        "Because of network latency, endpoint buffering, or offline hosts catching up on reconnection, logs may arrive at the ingestion server "
        "out of chronological order. To handle this, the server-side pipeline implements a three-timestamp model:"
    )
    p3.paragraph_format.line_spacing = 1.15
    p3.paragraph_format.space_after = Pt(8)

    bullet3_1 = doc.add_paragraph(style='List Bullet')
    run = bullet3_1.add_run("1. Event Time (event_time): ")
    run.bold = True
    bullet3_1.add_run("The UTC timestamp when the event actually occurred on the client endpoint. This is the source of truth for correlation.")
    bullet3_1.paragraph_format.space_after = Pt(3)

    bullet3_2 = doc.add_paragraph(style='List Bullet')
    run = bullet3_2.add_run("2. Ingestion Time (ingest_time): ")
    run.bold = True
    bullet3_2.add_run("The UTC timestamp when the Kafka broker received the record. Used to measure system ingestion latency.")
    bullet3_2.paragraph_format.space_after = Pt(3)

    bullet3_3 = doc.add_paragraph(style='List Bullet')
    run = bullet3_3.add_run("3. Processing Time (process_time): ")
    run.bold = True
    bullet3_3.add_run("The UTC timestamp when the database consumer parsed and wrote the record. Used for ingestion audits.")
    bullet3_3.paragraph_format.space_after = Pt(8)

    p3_handling = doc.add_paragraph(
        "Handling Strategies:\n"
        "• Database Ordering: The primary database uses event_time as its sorting and partitioning key, ensuring analyst queries are chronologically correct, regardless of arrival delays.\n"
        "• Watermarking / Windowing: Stream processing frameworks (like Flink or Kafka Streams) evaluate time-based correlations (e.g., detecting lateral movement within a 1-minute window) using Watermarks. Watermarks allow a predefined 'grace period' for late-arriving logs to ensure correlation logic triggers accurately."
    )
    p3_handling.paragraph_format.line_spacing = 1.15
    p3_handling.paragraph_format.left_indent = Inches(0.25)
    p3_handling.paragraph_format.space_after = Pt(12)

    # --- Question 4 ---
    add_heading_with_spacing(doc, "4) Storing Log Data Effectively", level=2)
    p4 = doc.add_paragraph(
        "Relational databases (OLTP) crash under the massive, sustained write rates of security event logging. "
        "To store SOC telemetry effectively, we design a tiered storage model utilizing a Columnar Time-Series Database (e.g., ClickHouse or TimescaleDB):"
    )
    p4.paragraph_format.line_spacing = 1.15
    p4.paragraph_format.space_after = Pt(8)

    bullet4_1 = doc.add_paragraph(style='List Bullet')
    run = bullet4_1.add_run("Columnar Storage: ")
    run.bold = True
    bullet4_1.add_run("Columnar engines group data on disk by column rather than row. This is ideal for security logs, which have many fields "
                           "but are queried on only a few columns at a time (e.g., 'Find all process names for IP 192.168.1.50'). This layout achieves "
                           "5x to 10x compression ratios, dramatically reducing hardware costs.")
    bullet4_1.paragraph_format.space_after = Pt(3)

    bullet4_2 = doc.add_paragraph(style='List Bullet')
    run = bullet4_2.add_run("Composite Indexing: ")
    run.bold = True
    bullet4_2.add_run("We define a primary composite index: `(event_time, endpoint_id, log_type)`. This optimizes typical security queries, "
                           "allowing the database engine to skip scanning unrelated files and retrieve data in milliseconds.")
    bullet4_2.paragraph_format.space_after = Pt(3)

    bullet4_3 = doc.add_paragraph(style='List Bullet')
    run = bullet4_3.add_run("Typed Fields: ")
    run.bold = True
    bullet4_3.add_run("Raw log strings are parsed into typed columns (e.g., IPv4/IPv6 data types for IPs, Low Cardinality enums for log types), "
                           "avoiding slow regular expression operations during analytics queries.")
    bullet4_3.paragraph_format.space_after = Pt(12)

    # --- Question 5 ---
    add_heading_with_spacing(doc, "5) Offset Management & Crash Recovery (Fault Tolerance)", level=2)
    p5 = doc.add_paragraph(
        "To ensure that a consumer recovers gracefully from crashes without losing logs or inserting duplicates (Exactly-Once Semantics), "
        "we implement manual offset commit management inside Apache Kafka combined with database idempotency constraints:"
    )
    p5.paragraph_format.line_spacing = 1.15
    p5.paragraph_format.space_after = Pt(8)

    bullet5_1 = doc.add_paragraph(style='List Bullet')
    run = bullet5_1.add_run("Disable Auto-Commit: ")
    run.bold = True
    bullet5_1.add_run("We set `enable.auto.commit = False` in the Kafka consumer configuration. This prevents Kafka from automatically "
                           "advancing the partition cursor before the logs are safely stored in the database.")
    bullet5_1.paragraph_format.space_after = Pt(3)

    bullet5_2 = doc.add_paragraph(style='List Bullet')
    run = bullet5_2.add_run("Post-Write Commit Cycle: ")
    run.bold = True
    bullet5_2.add_run("The consumer pulls a batch of messages, parses and enriches them, and writes them to the database. "
                           "Only after the database returns a successful write acknowledgment (ACK) does the consumer synchronously commit the offset "
                           "back to Kafka (`consumer.commit()`).")
    bullet5_2.paragraph_format.space_after = Pt(3)

    bullet5_3 = doc.add_paragraph(style='List Bullet')
    run = bullet5_3.add_run("Idempotent Writing: ")
    run.bold = True
    bullet5_3.add_run("If the consumer crashes after writing the logs to the database but before committing the offset to Kafka, the re-started consumer "
                           "will re-read and attempt to re-write the same batch. To prevent duplicate records, the database uses the log's unique `event_id` (UUIDv4) "
                           "as a unique constraint or primary key. On conflict, the database ignores or updates the existing entry, ensuring data integrity.")
    bullet5_3.paragraph_format.space_after = Pt(12)

    # --- Question 6 ---
    add_heading_with_spacing(doc, "6) Real-Time Formatting and In-Stream Enrichment (No Raw Storage)", level=2)
    p6 = doc.add_paragraph(
        "Storing raw, unparsed logs directly in a database is highly inefficient because it requires expensive runtime parsing during searches. "
        "Our architecture uses an in-stream parsing and enrichment model inside the consumer application before database ingestion:"
    )
    p6.paragraph_format.line_spacing = 1.15
    p6.paragraph_format.space_after = Pt(8)

    bullet6_1 = doc.add_paragraph(style='List Bullet')
    run = bullet6_1.add_run("Structured Schema Conversion: ")
    run.bold = True
    bullet6_1.add_run("Raw string values (e.g., syslog messages) are parsed using regex, grok filters, or schema definitions "
                           "and converted into structured objects.")
    bullet6_1.paragraph_format.space_after = Pt(3)

    bullet6_2 = doc.add_paragraph(style='List Bullet')
    run = bullet6_2.add_run("Memory-Cached GeoIP & WHOIS Lookups: ")
    run.bold = True
    bullet6_2.add_run("To enrich events without slow API calls, the consumer queries local memory-mapped databases (e.g., MaxMind MaxISP cache) "
                           "and persistent WHOIS cache tables to map IPs to locations/ASNs and domains to registration ages.")
    bullet6_2.paragraph_format.space_after = Pt(3)

    bullet6_3 = doc.add_paragraph(style='List Bullet')
    run = bullet6_3.add_run("High-Speed Threat Intelligence Matching: ")
    run.bold = True
    bullet6_3.add_run("Domains and IPs are matched against local blacklists and Threat Intelligence feeds stored in high-performance "
                           "in-memory databases (e.g., Redis) before ingestion.")
    bullet6_3.paragraph_format.space_after = Pt(12)

    # --- Question 7 ---
    add_heading_with_spacing(doc, "7) Retention Policy and Archival Storage (Tiered Storage)", level=2)
    p7 = doc.add_paragraph(
        "To manage costs while meeting long-term security compliance standards, we implement an Index Lifecycle Management (ILM) flow "
        "consisting of three distinct tiers:"
    )
    p7.paragraph_format.line_spacing = 1.15
    p7.paragraph_format.space_after = Pt(8)

    # Table for ILM
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Storage Tier'
    hdr_cells[1].text = 'Age Range'
    hdr_cells[2].text = 'Hardware Media'
    hdr_cells[3].text = 'Operational Performance'
    
    for cell in hdr_cells:
        set_cell_background(cell, "1F4E79")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    tiers = [
        ("Hot (Active)", "0 – 15 Days", "High-Speed SSDs (Primary DB)", "Sub-second analytical queries, active incident investigation."),
        ("Warm (Searchable)", "16 – 90 Days", "High-Capacity HDDs", "Compressed DB tables, slower threat hunting and trend mapping."),
        ("Cold (Archival)", "91+ Days", "S3 / MinIO Object Storage", "Exported as compressed columnar Parquet files. Queryable via S3 Athena/DuckDB.")
    ]

    for t_name, t_age, t_media, t_perf in tiers:
        row_cells = table.add_row().cells
        row_cells[0].text = t_name
        row_cells[1].text = t_age
        row_cells[2].text = t_media
        row_cells[3].text = t_perf
        for p in row_cells[0].paragraphs:
            for r in p.runs:
                r.font.bold = True

    p7_end = doc.add_paragraph()
    p7_end.paragraph_format.space_before = Pt(12)
    p7_end.paragraph_format.space_after = Pt(12)

    # --- Question 8 ---
    add_heading_with_spacing(doc, "8) Kafka Topic & Partition Design", level=2)
    p8 = doc.add_paragraph(
        "To manage high throughput and enable parallel processing while preserving the chronological order of endpoint events, "
        "we use the following Kafka Topic and Partition design:"
    )
    p8.paragraph_format.line_spacing = 1.15
    p8.paragraph_format.space_after = Pt(8)

    bullet8_1 = doc.add_paragraph(style='List Bullet')
    run = bullet8_1.add_run("Topic Allocation: ")
    run.bold = True
    bullet8_1.add_run("We define exactly two topics to separate typical data from high-priority security findings:\n"
                           "  • `dns-events-raw`: Handles high-volume, routine activity logs (benign traffic, typical resolutions) used for historical searches and baseline modeling.\n"
                           "  • `dns-alerts`: Dedicated to high-severity findings (DGA, Tunneling, Typosquatting, known malware indicators) requiring immediate ingestion, security analyst alerts, and incident response routing.")
    bullet8_1.paragraph_format.space_after = Pt(3)

    bullet8_2 = doc.add_paragraph(style='List Bullet')
    run = bullet8_2.add_run("Partition Count: ")
    run.bold = True
    bullet8_2.add_run("We configure each topic with 3 partitions. This supports horizontal scaling, allowing up to 3 consumer instances per consumer group to read and process events in parallel.")
    bullet8_2.paragraph_format.space_after = Pt(3)

    bullet8_3 = doc.add_paragraph(style='List Bullet')
    run = bullet8_3.add_run("Partitioning Key Strategy: ")
    run.bold = True
    bullet8_3.add_run("We use `endpoint_id` (or `hostname`) as the message routing key. Kafka hashes this key using MurmurHash2 and routes it:\n"
                           "  Partition = Hash(endpoint_id) % Total Partitions\n"
                           "This guarantees that all logs from a single endpoint map to the same partition and are processed sequentially. "
                           "This is crucial for timeline reconstruction during incident investigations.")
    bullet8_3.paragraph_format.space_after = Pt(24)

    # --- Question 9 / Diagram ---
    add_heading_with_spacing(doc, "9) Kafka Integration Telemetry Pipeline Flow Chart", level=1, space_before=18)
    
    p9 = doc.add_paragraph(
        "The following visual flowchart details the complete, end-to-end telemetry pipeline flow. It traces a security log event "
        "from host-level collection (Stage 1), routing and partition hashing inside Apache Kafka (Stage 2), parsing and in-stream "
        "enrichment inside the consumer group service (Stage 3), and finally, the database commit execution and long-term storage archival lifecycle (Stage 4)."
    )
    p9.paragraph_format.line_spacing = 1.15
    p9.paragraph_format.space_after = Pt(18)

    # Insert Image
    if os.path.exists("kafka_integration_diagram.png"):
        doc.add_picture("kafka_integration_diagram.png", width=Inches(6.0))
        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(6)
        p_cap.paragraph_format.space_after = Pt(12)
        run_cap = p_cap.add_run("Figure 9.1: End-to-End SOC Telemetry Ingestion, Brokering, Parsing, and Archival Flow Chart")
        run_cap.font.name = 'Arial'
        run_cap.font.size = Pt(9.5)
        run_cap.font.italic = True
        run_cap.font.color.rgb = RGBColor(128, 128, 128)
    else:
        doc.add_paragraph("[Error: kafka_integration_diagram.png not found. Diagram could not be embedded.]")

    # Save Document
    doc_name = "SOC_Telemetry_Kafka_Integration_Pipeline_Design.docx"
    doc.save(doc_name)
    print(f"Document generated successfully: {os.path.abspath(doc_name)}")

if __name__ == "__main__":
    main()
