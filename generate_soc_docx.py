import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_with_spacing(doc, text, level, space_before=12, space_after=6):
    heading = doc.add_heading(text, level=level)
    p_format = heading.paragraph_format
    p_format.space_before = Pt(space_before)
    p_format.space_after = Pt(space_after)
    p_format.keep_with_next = True
    return heading

def main():
    doc = Document()
    
    # --- Page Margins ---
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # --- Title ---
    title = doc.add_paragraph()
    title_format = title.paragraph_format
    title_format.space_before = Pt(12)
    title_format.space_after = Pt(4)
    title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("DeepCytes DNS Agent - SOC Visualization Blueprint")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(31, 78, 121)  # Deep Navy Blue

    # --- Subtitle ---
    subtitle = doc.add_paragraph()
    subtitle_format = subtitle.paragraph_format
    subtitle_format.space_after = Pt(18)
    subtitle_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run("Strategic Visual Analytics for Proactive DNS Telemetry Triage")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(128, 128, 128)

    # Divider line
    doc.add_paragraph("─" * 70).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Overview ---
    add_heading_with_spacing(doc, "1. Executive Overview", level=1, space_before=14)
    p_intro = doc.add_paragraph(
        "For a SOC (Security Operations Center) Analyst, raw DNS query logs represent a massive volume of telemetry. "
        "To quickly differentiate benign background traffic from active threat campaigns (like data exfiltration, "
        "malware beaconing, and phishing), visual representations are critical. This document outlines the key visual "
        "charts and frequency graphs implemented for the DNS Frontend Dashboard, detailing their utility and connection metrics."
    )
    p_intro.paragraph_format.line_spacing = 1.15
    p_intro.paragraph_format.space_after = Pt(10)

    # --- Visualizations Matrix ---
    add_heading_with_spacing(doc, "2. Dashboard Visualization Blueprint Matrix", level=1, space_before=14)
    
    # Create Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Chart Name / Type'
    hdr_cells[1].text = 'Visual Layout'
    hdr_cells[2].text = 'Operational Value for Analyst'
    hdr_cells[3].text = 'Telemetry Connection Points'
    
    # Set headers styling
    for cell in hdr_cells:
        set_cell_background(cell, "1F4E79")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(255, 255, 255)

    charts_info = [
        (
            "New Domain Detection Trend",
            "Frequency Line Graph (Hourly/Minute base)",
            "Identifies sudden bursts of newly-registered domains. An abnormal spike points to DGA malware beaconing, coordinated typosquatting drops, or phishing campaigns.",
            "Connects the domain creation timestamp from WHOIS parser with the current detection event time; flags 'is_newly_registered'."
        ),
        (
            "Log Severity Distribution",
            "Doughnut / Pie Chart",
            "Enables rapid alert prioritization and telemetry filtering. Shows the breakdown of Critical, Less Critical, and Informatory events at a glance.",
            "Groups log frequency based on risk scores: CRITICAL (Risk >= 60), LESS_CRITICAL (Risk 30-59), and INFORMATORY (Risk < 30)."
        ),
        (
            "Active Threat Vectors",
            "Horizontal Bar Chart",
            "Highlights which malicious activity vectors are currently dominant on the network. Allows analysts to spot active exfiltration vs. passive typosquatting.",
            "Counts occurrences of active rule triggers inside the 'alerts' array (e.g. Tunneling, DGA, Typosquatting, Threat Intel, Coruna)."
        ),
        (
            "Top Talkers / Domain Frequency",
            "Horizontal Bar Chart",
            "Pinpoints the top destination domains being queried. Helps identify anomalous persistent connections or high-frequency beacons.",
            "Connects 'dns.query' strings, counts query frequency, and tracks the initiating process metadata (Process Name/PID)."
        ),
        (
            "Entropy vs. Risk Score",
            "2D Scatter / Bubble Plot",
            "Visualizes outliers. High-entropy subdomains with high risk scores immediately cluster in the top-right quadrant, exposing covert tunneling activity.",
            "Connects character-based Shannon Entropy values of subdomains ('domain_analysis.entropy') with calculated 'risk.score'."
        )
    ]

    for name, layout, val, conn in charts_info:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = layout
        row_cells[2].text = val
        row_cells[3].text = conn
        # Style row headers
        for p in row_cells[0].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9.5)
        for i in range(1, 4):
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)

    # Add space after table
    doc.add_paragraph().paragraph_format.space_before = Pt(12)

    # --- Operational Triage Workflow ---
    add_heading_with_spacing(doc, "3. Analyst Triage Workflow integration", level=1, space_before=14)
    p_triage = doc.add_paragraph(
        "By integrating these visual charts with the endpoint's Process Lineage layer (command line, PID, and executable signer), "
        "a SOC analyst can quickly determine context. For example, if the 'New Domain Trend' line graph spikes: \n"
        "  1. The analyst checks the 'Active Threat Vectors' bar chart to see if it is typosquatting or tunneling.\n"
        "  2. They inspect the 'Severity Distribution' to find associated Critical events.\n"
        "  3. By clicking the event in the Live Stream table, they extract the originating Process (e.g. powershell.exe) and the Parent PID to block the compromised endpoint."
    )
    p_triage.paragraph_format.line_spacing = 1.15
    p_triage.paragraph_format.space_after = Pt(10)

    # Save Document
    doc_path = "DNS_SOC_Analyst_Visualizations.docx"
    doc.save(doc_path)
    print(f"SOC report generated successfully: {os.path.abspath(doc_path)}")

if __name__ == "__main__":
    main()
